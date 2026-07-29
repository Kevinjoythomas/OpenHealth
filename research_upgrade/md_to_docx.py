import re, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
src='research_upgrade/OpenHealth_Research_Paper_vNext.md'
dst='OpenHealth_Research_Paper_vNext.docx'
lines=open(src,encoding='utf-8').read().split('\n')
doc=Document()
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)

def add_runs(p, text):
    # split on **bold** and *italic*
    for seg in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            r=p.add_run(seg[2:-2]); r.bold=True
        elif seg.startswith('*') and seg.endswith('*'):
            r=p.add_run(seg[1:-1]); r.italic=True
        else:
            p.add_run(seg)

i=0
while i<len(lines):
    ln=lines[i]
    if ln.strip()=='':
        i+=1; continue
    # table block
    if ln.lstrip().startswith('|'):
        block=[]
        while i<len(lines) and lines[i].lstrip().startswith('|'):
            block.append(lines[i]); i+=1
        rows=[[c.strip() for c in r.strip().strip('|').split('|')] for r in block]
        rows=[r for r in rows if not all(set(c)<=set('-: ') for c in r)]  # drop separator
        if not rows: continue
        ncol=max(len(r) for r in rows)
        t=doc.add_table(rows=0, cols=ncol); t.style='Light Grid Accent 1'
        for ri,r in enumerate(rows):
            cells=t.add_row().cells
            for ci in range(ncol):
                txt=r[ci] if ci<len(r) else ''
                cell=cells[ci]; cell.text=''
                pp=cell.paragraphs[0]; add_runs(pp, txt)
                if ri==0:
                    for rn in pp.runs: rn.bold=True
        doc.add_paragraph('')
        continue
    if ln.startswith('### '):
        doc.add_heading(ln[4:].strip(), level=3)
    elif ln.startswith('## '):
        doc.add_heading(ln[3:].strip(), level=2)
    elif ln.startswith('# '):
        h=doc.add_heading(ln[2:].strip(), level=0)
    elif ln.startswith('> '):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Pt(18)
        add_runs(p, ln[2:].strip()); 
        for r in p.runs: r.italic=True
    elif re.match(r'^\s*-\s+', ln):
        p=doc.add_paragraph(style='List Bullet'); add_runs(p, re.sub(r'^\s*-\s+','',ln))
    elif re.match(r'^\d+\.\s', ln):
        p=doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\d+\.\s','',ln))
    else:
        p=doc.add_paragraph(); add_runs(p, ln)
    i+=1
doc.save(dst)
print("WROTE", dst)
